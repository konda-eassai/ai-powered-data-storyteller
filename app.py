# ============================================================
# dashboard/app.py
# Streamlit dashboard for HR Attrition Prediction
# ============================================================

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import joblib
from sklearn.pipeline import Pipeline

# ------------------------------------------------------------
# 1) Page config & custom CSS
# ------------------------------------------------------------
st.set_page_config(
    page_title="HR Attrition Predictor",
    page_icon="💼",
    layout="wide"
)

st.markdown(
    """
    <style>
    .title {font-size:2.3rem;font-weight:bold;color:#2A7E8C;}
    div.stButton>button {background-color:#2A7E8C;color:white;font-weight:bold;border-radius:6px;}
    [data-testid="stSidebar"] h1 {font-size:1.3rem;color:#2A7E8C;}
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown('<p class="title"> HR Analytics Employee Attrition Dashboard</p>',
            unsafe_allow_html=True)

# ------------------------------------------------------------
# 2) Load model & metrics
# ------------------------------------------------------------
MODEL_PATH = Path("../models/tuned_rf_model.pkl")
METRICS_PATH = Path("../models/test_metrics.pkl")

@st.cache_resource
def load_model():
    return joblib.load(MODEL_PATH)

@st.cache_data
def load_metrics():
    return joblib.load(METRICS_PATH)

model: Pipeline = load_model()
metrics = load_metrics()

# Build feature importance info
numeric_features = model.named_steps["preprocessor"].transformers_[0][2]
cat_encoder = model.named_steps["preprocessor"].named_transformers_["cat"]
cat_feats = cat_encoder.feature_names_in_
onehot_cols = cat_encoder.get_feature_names_out(cat_feats)
all_features = list(numeric_features) + list(onehot_cols)

importances = model.named_steps["classifier"].feature_importances_
feat_imp = pd.Series(importances, index=all_features).sort_values(ascending=False)

# ------------------------------------------------------------
# 3) Layout: Metrics + Features side-by-side
# ------------------------------------------------------------
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("### Top Features Driving Attrition")
    fig, ax = plt.subplots(figsize=(9,6))
    sns.barplot(x=feat_imp.values[:15], y=feat_imp.index[:15], palette="viridis", ax=ax)
    ax.set_title("Top 15 Features")
    st.pyplot(fig)

with col2:
    st.markdown("### Model Performance")
    st.dataframe(pd.DataFrame(metrics, index=["Score"]).T.style.format("{:.3f}"))

st.markdown("---")

# ------------------------------------------------------------
# 4) Tabs for EDA / Predictions / Insights
# ------------------------------------------------------------
tab1, tab2, tab3 = st.tabs(["EDA", "Predictions", "Model Insights"])

# ---------- Tab 1: EDA ----------
with tab1:
    st.subheader("Explore Your Dataset")
    uploaded_file = st.file_uploader("Upload a CSV file", type=["csv"], key="eda")

    if uploaded_file:
        user_df = pd.read_csv(uploaded_file)
        st.success(f"{user_df.shape[0]} rows × {user_df.shape[1]} columns loaded.")
        st.dataframe(user_df.head())
        st.write("Missing values:", user_df.isna().sum())
        st.write(user_df.describe(include='all').T.head(10))

        if st.checkbox("Show detailed plots"):
            num_cols = user_df.select_dtypes(include=["int64","float64"]).columns
            if len(num_cols) > 0:
                user_df[num_cols].hist(figsize=(12,8), bins=20, edgecolor="black")
                st.pyplot(plt.gcf())

            if len(num_cols) > 1:
                fig, ax = plt.subplots(figsize=(8,6))
                sns.heatmap(user_df[num_cols].corr(), annot=True, cmap="coolwarm", fmt=".2f", ax=ax)
                st.pyplot(fig)

# ---------- Tab 2: Predictions ----------
with tab2:
    st.subheader("Predict Attrition")
    pred_file = st.file_uploader("Upload HR-style CSV for prediction", type=["csv"], key="pred")

    if pred_file:
        df_pred = pd.read_csv(pred_file)
        if all(c in df_pred.columns for c in model.feature_names_in_):
            probs = model.predict_proba(df_pred[model.feature_names_in_])[:, 1]
            preds = ["Yes" if p > 0.5 else "No" for p in probs]
            results = pd.DataFrame({"Probability": probs, "Prediction": preds})
            st.dataframe(results.head(20))
            st.download_button(
                "Download Predictions",
                results.to_csv(index=False).encode("utf-8"),
                "predictions.csv",
                "text/csv"
            )
        else:
            st.warning("Uploaded file doesn’t match the model’s schema.")

# ---------- Tab 3: Model Insights ----------
with tab3:
    st.subheader("Model Insights")

    if st.checkbox("Show Top Features Driving Attrition", value=True):
        fig2, ax = plt.subplots(figsize=(9, 6))
        sns.barplot(
            x=feat_imp.values[:15],
            y=feat_imp.index[:15],
            palette="viridis",
            ax=ax
        )
        ax.set_title("Top 15 Drivers of Attrition")
        st.pyplot(fig2)


    st.markdown("### Test Metrics")
    st.dataframe(pd.DataFrame(metrics, index=["Score"]).T.style.format("{:.3f}"))

st.markdown("---")
st.caption("© 2025 Built by Konda Eassai | Imarticus Data Science Internship")


from fpdf import FPDF
import io

# --- Button to download PDF report ---
if st.button("📄 Download Full Report (PDF)"):
    # 1) Save current feature importance figure to bytes
    img_bytes = io.BytesIO()
    fig2.savefig(img_bytes, format="png", bbox_inches="tight")
    img_bytes.seek(0)

    # 2) Build PDF
    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "HR Attrition Model Report", ln=True, align="C")

    pdf.ln(8)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 8,
        "This report summarizes the tuned RandomForest model for predicting employee attrition.\n\n"
        "Below are the evaluation metrics on the test set:"
    )

    pdf.ln(4)
    for k, v in metrics.items():
        pdf.cell(0, 8, f"{k}: {v:.3f}", ln=True)

    pdf.ln(6)
    pdf.multi_cell(0, 8, "Top 15 features contributing to attrition are shown below:")

    # 3) Insert the plot
    pdf.ln(4)
    pdf.image(img_bytes, x=None, y=None, w=170)

    pdf.ln(10)
    pdf.set_font("Arial", "I", 10)
    pdf.cell(0, 8, "Report generated by Streamlit Dashboard", ln=True, align="C")

    # 4) Output PDF as bytes
    pdf_bytes = pdf.output(dest="S").encode("latin1") if isinstance(pdf.output(dest="S"), str) else pdf.output(dest="S")

    st.download_button(
        label="💾 Save Report as PDF",
        data=bytes(pdf_bytes),
        file_name="HR_Attrition_Report.pdf",
        mime="application/pdf"
    )


from docx import Document
import os

# --- Button to download Word report ---
if st.button("📝 Download Full Report (Word)"):
    # 1) Create a Word document
    doc = Document()
    doc.add_heading("HR Attrition Model Report", 0)
    doc.add_paragraph(
        "This report summarizes the tuned RandomForest model for predicting employee attrition.\n\n"
        "Below are the evaluation metrics on the test set:"
    )

    for k, v in metrics.items():
        doc.add_paragraph(f"{k}: {v:.3f}")

    doc.add_paragraph("\nTop 15 features contributing to attrition are shown below:")

    # 2) Save the feature-importance figure temporarily
    img_path = "feat_importance_tmp.png"
    fig2.savefig(img_path, bbox_inches="tight")

    doc.add_picture(img_path, width=None)

    if os.path.exists(img_path):
        os.remove(img_path)

    doc.add_paragraph("\nReport generated by Streamlit Dashboard.")

    # 3) Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    st.download_button(
        label="💾 Save Report as Word",
        data=bio.getvalue(),
        file_name="HR_Attrition_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
