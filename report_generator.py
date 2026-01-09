# dashboard/report_generator.py
from fpdf import FPDF
from docx import Document
import io
import os
from pathlib import Path
import pandas as pd

def make_pdf_report(metrics: dict, fig, out_path="reports/hr_report.pdf"):
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    img_bytes = io.BytesIO()
    fig.savefig(img_bytes, format="png", bbox_inches='tight')
    img_bytes.seek(0)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "HR Attrition Model Report", ln=True, align="C")
    pdf.ln(8)
    pdf.set_font("Arial", "", 12)
    for k, v in metrics.items():
        pdf.cell(0, 8, f"{k}: {v:.3f}", ln=True)
    pdf.ln(6)
    pdf.image(img_bytes, x=10, w=190)
    pdf.output(out_path)
    return out_path

def make_word_report(metrics: dict, fig, out_path="reports/hr_report.docx"):
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    img_path = "tmp_feat.png"
    fig.savefig(img_path, bbox_inches='tight')
    doc = Document()
    doc.add_heading("HR Attrition Model Report", 0)
    for k, v in metrics.items():
        doc.add_paragraph(f"{k}: {v:.3f}")
    doc.add_picture(img_path, width=None)
    doc.save(out_path)
    if os.path.exists(img_path):
        os.remove(img_path)
    return out_path
